import mammoth
import openpyxl
import pypandoc
from docx import Document

from Natura_config import data_dir


def convert_template(file, format):
    file = data_dir + "\\" + file
    if format == 'md':
        with open(file, "rb") as docx_file:
            result = mammoth.convert_to_markdown(docx_file)
    if format == 'html':
        with open(file, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
    file = file.replace('docx', '')
    with open(file + format, "w", encoding="utf-8") as new_file:
        new_file.write(result.value)


def create_report_md(template, new_file_name):
    filename = data_dir + "\\" + template
    new_file = data_dir + "\\" + new_file_name
    table = data_dir + "\\Donja Posavina_0.xlsx.md"
    orig = 'Table\_1'
    with open(table, 'r') as f:
        table_in_md_string = f.read()
    with open(filename, 'r') as f:
        template_text = f.read()
        new_text = template_text.replace(orig, table_in_md_string)
    with open(new_file, 'w') as f:
        f.write(new_text)
        print("wrote")
        print(new_file)

def create_report_html(template, new_file_name):
    filename = data_dir + "\\" + template
    new_file = data_dir + "\\" + new_file_name
    table = data_dir + "\\Donja Posavina_0.xlsx.html"
    orig = 'Table_1'
    with open(table, 'r') as f:
        new = f.readlines()
    with open(filename, 'r') as f:
        text = f.readlines()
    for line in text:
        new_text = line.replace(orig, ''.join(new))
    with open(new_file, 'w') as f:
        f.write(new_text)
        print("wrote")
        print(new_file)

def report_knit(file_to_knit, new_file_name):
    file_to_knit = data_dir + "\\" + file_to_knit
    new_file = data_dir + "\\" + new_file_name
    pypandoc.convert_file(file_to_knit, to='docx', outputfile=new_file)

def insert_excel():
    workbook = openpyxl.load_workbook(data_dir + "\\" + 'Donja Posavina_0.xlsx')
    worksheet = workbook['Sheet1']
    document = Document(data_dir + "\\" + 'Template.docx')

    # Add a table to the document and fill it with data from the Excel sheet
    for paragraph in document.paragraphs:
        if 'Table_1' in paragraph.text:
            table = document.add_table(rows=1, cols=len(worksheet[1]))
            table.style = 'DV_tablica'

            # Add the column names from the first row of the Excel sheet
            for i, cell in enumerate(worksheet[1]):
                table.cell(0, i).text = cell.value

            # Add the data from the rest of the rows in the Excel sheet
            for row in worksheet.iter_rows(min_row=2):
                cells = []
                for cell in row:
                    cells.append(cell.value)
                row_cells = table.add_row().cells
                for i, cell in enumerate(row_cells):
                    cell.text = str(cells[i])
            paragraph.text = paragraph.text.replace('Table_1', '')
            new_paragraph = document.add_paragraph('')
            #table.insert_paragraph_before(new_paragraph)
            #table.add_paragraph(new_paragraph)

    # Save the docx document
    document.save('report.docx')


def obtain_styles():
    # Open the source docx file
    source_doc = Document(data_dir + "\\"+'Tekst_SPUO_IDPPZ_Medj_zup-priroda_NG_es_v1.docx')

    # Open the destination docx file
    destination_doc = Document(data_dir + "\\" + 'New_document.docx')

    # Iterate through all styles in the source document
    for style in source_doc.styles:
        # Check if the style is not a default style
        if not style.builtin:
            # Create a new style in the destination document with the same name and attributes as the source style
            destination_doc.styles.add_style(style.name, style)

    # Save the destination document
    destination_doc.save(data_dir + "\\" + 'destination.docx')

